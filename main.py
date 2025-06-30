from datetime import date
from io import BytesIO
from typing import List, Optional, Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field, EmailStr

app = FastAPI(
    title="CV Generator API",
    description="API to generate CV Word documents from JSON data",
    version="0.1.0"
)


# Pydantic models based on the provided JSON schema
class WorkExperience(BaseModel):
    position: str
    company_name: str
    location: str
    responsibilities: str
    period: str


class Education(BaseModel):
    institution: str
    period_from: int
    period_to: int
    specialization: str
    location: str
    location_of_graduation: str


class LanguageProficiency(BaseModel):
    language: str
    writing: int = Field(ge=2, le=5)
    speaking: int = Field(ge=2, le=5)
    understanding: int = Field(ge=2, le=5)


class Residency(BaseModel):
    city: str
    country: str


class CV(BaseModel):
    full_name: str
    phone_number: str
    phone_number_2: Optional[str] = None
    email: EmailStr
    nationality: Optional[str] = None
    position: Optional[str] = None
    employment: Optional[Literal["looking", "employed", "freelance"]] = None
    location: Optional[str] = None
    skills: Optional[List[str]] = None
    work_experience: Optional[List[WorkExperience]] = None
    education: Optional[List[Education]] = None
    language_proficiency: Optional[List[LanguageProficiency]] = None
    marriage_status: Optional[str] = None
    have_children: Optional[Literal["Yes", "No"]] = None
    date_of_birth: Optional[date] = None
    place_of_birth: Optional[str] = None
    residency: Optional[Residency] = None
    age: Optional[int] = Field(None, ge=0)
    height: Optional[int] = Field(None, ge=30, description="Height in centimeters")
    weight: Optional[int] = Field(None, ge=2, description="Weight in kilograms")


def json_to_docx(cv_data: CV) -> BytesIO:
    """Convert CV JSON data directly to Word document format"""
    doc = Document()

    # Set up document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title - Full Name
    title = doc.add_heading(cv_data.full_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Information
    doc.add_heading('Contact Information', level=1)

    # Create a table for contact info and photo
    contact_table = doc.add_table(rows=1, cols=2)
    contact_table.autofit = False

    # Set column widths (70% for contact info, 30% for photo)
    contact_table.columns[0].width = Inches(4.5)  # Adjust as needed
    contact_table.columns[1].width = Inches(1.5)  # Adjust as needed

    # Left cell for contact information
    left_cell = contact_table.cell(0, 0)
    contact_info = left_cell.paragraphs[0]
    contact_info.add_run(f'Phone: {cv_data.phone_number}').bold = True
    contact_info.add_run('\n')
    if cv_data.phone_number_2:
        contact_info.add_run(f'Secondary Phone: {cv_data.phone_number_2}').bold = True
        contact_info.add_run('\n')
    contact_info.add_run(f'Email: {cv_data.email}').bold = True
    contact_info.add_run('\n')
    if cv_data.location:
        contact_info.add_run(f'Location: {cv_data.location}').bold = True

    # Right cell for photo placeholder
    right_cell = contact_table.cell(0, 1)
    photo_info = right_cell.paragraphs[0]
    photo_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_info.add_run('Photo Placeholder').italic = True

    # Add a blank paragraph after the table
    doc.add_paragraph()

    # Personal Information
    personal_info_items = []
    if cv_data.nationality:
        personal_info_items.append(f'Nationality: {cv_data.nationality}')
    if cv_data.date_of_birth:
        personal_info_items.append(f'Date of Birth: {cv_data.date_of_birth}')
    if cv_data.place_of_birth:
        personal_info_items.append(f'Place of Birth: {cv_data.place_of_birth}')
    if cv_data.age:
        personal_info_items.append(f'Age: {cv_data.age}')
    if cv_data.marriage_status:
        personal_info_items.append(f'Marriage Status: {cv_data.marriage_status}')
    if cv_data.have_children:
        personal_info_items.append(f'Have Children: {cv_data.have_children}')
    if cv_data.height:
        personal_info_items.append(f'Height: {cv_data.height} cm')
    if cv_data.weight:
        personal_info_items.append(f'Weight: {cv_data.weight} kg')
    if cv_data.residency:
        personal_info_items.append(f'Residency: {cv_data.residency.city}, {cv_data.residency.country}')

    if personal_info_items:
        doc.add_heading('Personal Information', level=1)
        personal_info = doc.add_paragraph()
        for item in personal_info_items:
            personal_info.add_run(item).bold = True
            personal_info.add_run('\n')

    # Professional Information
    if cv_data.position:
        doc.add_heading('Professional Summary', level=1)
        prof_info = doc.add_paragraph()
        prof_info.add_run(f'Position: {cv_data.position}').bold = True
        prof_info.add_run('\n')
        if cv_data.employment:
            prof_info.add_run(f'Employment Status: {cv_data.employment.capitalize()}').bold = True

    # Skills
    if cv_data.skills:
        doc.add_heading('Skills', level=1)
        skills_para = doc.add_paragraph()
        for skill in cv_data.skills:
            skills_para.add_run(f'• {skill}\n')

    # Education
    if cv_data.education:
        doc.add_heading('Education', level=1)
        for edu in cv_data.education:
            doc.add_heading(f'{edu.institution}', level=2)
            edu_para = doc.add_paragraph()
            edu_para.add_run('Specialization: ').bold = True
            edu_para.add_run(f'{edu.specialization}\n')
            edu_para.add_run('Period: ').bold = True
            edu_para.add_run(f'{edu.period_from} - {edu.period_to}\n')
            edu_para.add_run('Location: ').bold = True
            edu_para.add_run(f'{edu.location}\n')
            edu_para.add_run('Location of Graduation: ').bold = True
            edu_para.add_run(f'{edu.location_of_graduation}')
            doc.add_paragraph()

    # Language Proficiency
    if cv_data.language_proficiency:
        doc.add_heading('Language Proficiency', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Language'
        header_cells[1].text = 'Writing'
        header_cells[2].text = 'Speaking'
        header_cells[3].text = 'Understanding'

        # Add language data
        for lang in cv_data.language_proficiency:
            row_cells = table.add_row().cells
            row_cells[0].text = lang.language
            row_cells[1].text = f'{lang.writing}'
            row_cells[2].text = f'{lang.speaking}'
            row_cells[3].text = f'{lang.understanding}'

    # Work Experience
    if cv_data.work_experience:
        doc.add_heading('Work Experience', level=1)
        for job in cv_data.work_experience:
            doc.add_heading(f'{job.position} at {job.company_name}', level=2)
            job_para = doc.add_paragraph()
            job_para.add_run('Location: ').bold = True
            job_para.add_run(f'{job.location}\n')
            job_para.add_run('Period: ').bold = True
            job_para.add_run(f'{job.period}\n')
            job_para.add_run('Responsibilities:\n').bold = True
            job_para.add_run(f'{job.responsibilities}')
            doc.add_paragraph()

    # Save document to memory
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes


@app.get("/")
async def root():
    return {
        "message": "CV Generator API",
        "usage": "POST your CV data to /generate-cv endpoint to get a Word document"
    }


@app.post("/generate-cv", response_class=Response)
async def generate_cv(cv_data: CV):
    try:
        # Convert JSON directly to Word document
        docx_bytes = json_to_docx(cv_data)

        # Return the Word document
        return Response(
            content=docx_bytes.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={cv_data.full_name.replace(' ', '_')}_CV.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating CV: {str(e)}")
